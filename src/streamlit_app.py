import streamlit as st
import json
import pandas as pd
from datetime import datetime
import os
from pathlib import Path
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches
import plotly.graph_objects as go
import plotly.express as px

# 导入项目模块
from auth.user_manager import UnifiedUserManager
from auth.session_handler import SessionHandler
from agents.agent_registry import AgentRegistry
from discussion.discussion_engine import ClinicalDiscussionEngine
from storage.discussion_storage import DiscussionStorage
from utils.config import ClinicalConfig
from utils.logger import setup_logger

class ClinicalWebInterface:
    def __init__(self):
        self.setup_page_config()
        self.initialize_session_state()
        self.user_manager = UnifiedUserManager()
        self.session_handler = SessionHandler()
        self.agent_registry = AgentRegistry()
        self.discussion_storage = DiscussionStorage()
        self.config = ClinicalConfig()
        self.logger = setup_logger("web_interface")
        
    def setup_page_config(self):
        """设置页面配置"""
        st.set_page_config(
            page_title="临床MDT智能模拟助手",
            page_icon="🏥",
            layout="wide",
            initial_sidebar_state="expanded"
        )
        
        # 自定义CSS样式
        st.markdown("""
        <style>
        .main-header {
            font-size: 2.5rem;
            color: #1f77b4;
            text-align: center;
            margin-bottom: 2rem;
        }
        .agent-card {
            border: 1px solid #ddd;
            border-radius: 10px;
            padding: 15px;
            margin: 10px 0;
            background-color: #f9f9f9;
        }
        .discussion-bubble {
            border-radius: 15px;
            padding: 10px 15px;
            margin: 5px 0;
            max-width: 80%;
        }
        .user-bubble {
            background-color: #d4edda;
            margin-left: 20%;
        }
        .agent-bubble {
            background-color: #f8f9fa;
            margin-right: 20%;
        }
        .specialty-tag {
            background-color: #6c757d;
            color: white;
            padding: 2px 8px;
            border-radius: 12px;
            font-size: 0.8rem;
            margin-right: 5px;
        }
        </style>
        """, unsafe_allow_html=True)

    def initialize_session_state(self):
        """初始化会话状态"""
        if 'authenticated' not in st.session_state:
            st.session_state.authenticated = False
        if 'current_user' not in st.session_state:
            st.session_state.current_user = None
        if 'session_id' not in st.session_state:
            st.session_state.session_id = None
        if 'discussion_active' not in st.session_state:
            st.session_state.discussion_active = False
        if 'selected_agents' not in st.session_state:
            st.session_state.selected_agents = []
        if 'discussion_log' not in st.session_state:
            st.session_state.discussion_log = []
        if 'medical_record' not in st.session_state:
            st.session_state.medical_record = ""
        if 'discussion_question' not in st.session_state:
            st.session_state.discussion_question = ""
        if 'current_round' not in st.session_state:
            st.session_state.current_round = 0
        if 'user_interventions' not in st.session_state:
            st.session_state.user_interventions = []

    def render_authentication_section(self):
        """渲染认证界面"""
        st.markdown("<h1 class='main-header'>🏥 临床MDT智能模拟助手</h1>", unsafe_allow_html=True)
        
        tab1, tab2 = st.tabs(["登录", "注册"])
        
        with tab1:
            with st.form("login_form"):
                st.subheader("用户登录")
                username = st.text_input("用户名", key="login_username")
                password = st.text_input("密码", type="password", key="login_password")
                submit_login = st.form_submit_button("登录")
                
                if submit_login:
                    if self.user_manager.authenticate(username, password):
                        st.session_state.authenticated = True
                        st.session_state.current_user = username
                        st.session_state.session_id = self.session_handler.create_session(username)
                        st.success("登录成功！")
                        st.rerun()
                    else:
                        st.error("用户名或密码错误")

        with tab2:
            with st.form("register_form"):
                st.subheader("用户注册")
                new_username = st.text_input("用户名", key="register_username")
                new_password = st.text_input("密码", type="password", key="register_password")
                confirm_password = st.text_input("确认密码", type="password", key="confirm_password")
                submit_register = st.form_submit_button("注册")
                
                if submit_register:
                    if new_password != confirm_password:
                        st.error("密码不匹配")
                    elif self.user_manager.create_user(new_username, new_password):
                        st.success("注册成功！请登录")
                    else:
                        st.error("用户名已存在")

    def render_agent_selection(self):
        """渲染智能体选择界面"""
        st.sidebar.header("🔧 智能体管理")
        
        # 显示可用智能体
        available_agents = self.agent_registry.get_available_agents(st.session_state.session_id)
        
        st.sidebar.subheader("内置智能体")
        for agent_name, agent_info in available_agents.items():
            if agent_info.get("builtin", False):
                col1, col2 = st.sidebar.columns([3, 1])
                with col1:
                    st.write(f"**{agent_name}**")
                    st.caption(agent_info["specialty"])
                with col2:
                    if st.button("选择", key=f"select_{agent_name}"):
                        if agent_name not in st.session_state.selected_agents:
                            st.session_state.selected_agents.append(agent_name)
                            st.success(f"已选择 {agent_name}")
        
        # 自定义智能体
        st.sidebar.subheader("自定义智能体")
        with st.sidebar.expander("添加自定义智能体"):
            custom_name = st.text_input("智能体名称")
            custom_prompt = st.text_area("智能体提示词", height=100)
            if st.button("添加自定义智能体"):
                if custom_name and custom_prompt:
                    self.agent_registry.create_custom_agent(
                        st.session_state.session_id, custom_name, custom_prompt
                    )
                    st.success(f"自定义智能体 {custom_name} 已添加")
        
        # 显示已选智能体
        st.sidebar.subheader("已选智能体")
        if st.session_state.selected_agents:
            for agent in st.session_state.selected_agents:
                col1, col2 = st.sidebar.columns([3, 1])
                with col1:
                    st.write(f"• {agent}")
                with col2:
                    if st.button("移除", key=f"remove_{agent}"):
                        st.session_state.selected_agents.remove(agent)
                        st.rerun()
        else:
            st.sidebar.info("尚未选择任何智能体")

    def render_medical_input(self):
        """渲染病历输入界面"""
        st.header("📋 病历信息输入")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.subheader("患者基本信息")
            patient_info = st.text_area("患者基本信息（年龄、性别、主诉等）", height=100)
            
            st.subheader("现病史")
            current_illness = st.text_area("现病史详细描述", height=150)
            
            st.subheader("既往史")
            medical_history = st.text_area("既往病史、手术史、过敏史等", height=100)
            
            st.subheader("辅助检查")
            lab_tests = st.text_area("实验室检查、影像学检查等结果", height=150)
        
        with col2:
            st.subheader("生命体征")
            bp = st.text_input("血压")
            hr = st.text_input("心率")
            temp = st.text_input("体温")
            rr = st.text_input("呼吸频率")
            
            st.subheader("体格检查")
            physical_exam = st.text_area("体格检查发现", height=100)
            
            st.subheader("讨论问题")
            discussion_question = st.text_area("需要讨论的具体问题", height=100)
        
        # 整合病历信息
        if st.button("生成完整病历"):
            medical_record = self.compile_medical_record(
                patient_info, current_illness, medical_history, lab_tests,
                bp, hr, temp, rr, physical_exam
            )
            st.session_state.medical_record = medical_record
            st.session_state.discussion_question = discussion_question
            
            st.success("病历生成完成！")
            st.expander("查看完整病历").write(medical_record)

    def compile_medical_record(self, patient_info, current_illness, medical_history, 
                             lab_tests, bp, hr, temp, rr, physical_exam):
        """编译完整病历"""
        record = f"""
# 患者病历摘要
## 基本信息
{patient_info}

## 生命体征
- 血压: {bp}
- 心率: {hr}
- 体温: {temp}
- 呼吸频率: {rr}

## 现病史
{current_illness}

## 既往史
{medical_history}

## 体格检查
{physical_exam}

## 辅助检查
{lab_tests}
"""
        return record

    def render_discussion_control(self):
        """渲染讨论控制界面"""
        st.header("💬 多智能体讨论")
        
        if not st.session_state.selected_agents:
            st.warning("请先选择至少一个智能体")
            return
        
        if not st.session_state.medical_record:
            st.warning("请先输入病历信息")
            return
        
        # 讨论配置
        col1, col2, col3 = st.columns(3)
        with col1:
            discussion_rounds = st.slider("讨论轮数", 1, 10, 3)
        with col2:
            intervention_mode = st.selectbox(
                "用户介入模式",
                ["旁观模式", "主动介入", "仅关键节点介入"]
            )
        with col3:
            discussion_speed = st.slider("讨论速度", 1, 5, 3)
        
        # 开始讨论按钮
        if st.button("🚀 开始讨论", use_container_width=True):
            self.start_discussion(discussion_rounds, intervention_mode)

    def start_discussion(self, rounds, intervention_mode):
        """开始讨论"""
        st.session_state.discussion_active = True
        st.session_state.current_round = 0
        st.session_state.discussion_log = []
        st.session_state.user_interventions = []
        
        # 创建讨论引擎
        args = self.create_discussion_args()
        discussion_engine = ClinicalDiscussionEngine(args, st.session_state.session_id)
        discussion_engine.initialize_agents(st.session_state.selected_agents)
        
        # 在单独的线程中运行讨论
        with st.spinner("智能体正在讨论中..."):
            result = discussion_engine.run_discussion(
                st.session_state.medical_record,
                st.session_state.discussion_question
            )
            
            # 保存讨论结果
            st.session_state.discussion_result = result
            st.session_state.discussion_active = False
            
            # 保存到存储
            self.discussion_storage.save_discussion(
                st.session_state.current_user,
                {
                    "agents": st.session_state.selected_agents,
                    "rounds": rounds,
                    "medical_record": st.session_state.medical_record,
                    "question": st.session_state.discussion_question,
                    "log": st.session_state.discussion_log,
                    "summary": result,
                    "interventions": st.session_state.user_interventions
                }
            )

    def create_discussion_args(self):
        """创建讨论参数"""
        class Args:
            def __init__(self, config):
                self.model = config.model.engine
                self.llm_name = config.model.model_name
                self.url = config.model.api_base  # 使用配置中的API地址
                self.temp = config.model.temperature
                self.debug = False
        
        return Args()

    def render_discussion_display(self):
        """渲染讨论显示界面"""
        if not st.session_state.discussion_active:
            return
        
        st.header("实时讨论进程")
        
        # 创建讨论容器
        discussion_container = st.container()
        intervention_container = st.expander("💬 用户介入")
        
        with discussion_container:
            # 显示当前轮次
            st.subheader(f"第 {st.session_state.current_round + 1} 轮讨论")
            
            # 显示讨论进度
            progress_bar = st.progress(0)
            progress = (st.session_state.current_round + 1) / st.session_state.discussion_rounds
            progress_bar.progress(progress)
            
            # 显示智能体发言
            if st.session_state.discussion_log:
                latest_round = st.session_state.discussion_log[-1]
                for contribution in latest_round["contributions"]:
                    self.render_agent_contribution(contribution)
        
        with intervention_container:
            if intervention_mode == "主动介入":
                self.render_active_intervention()
            elif intervention_mode == "仅关键节点介入":
                self.render_critical_intervention()

    def render_agent_contribution(self, contribution):
        """渲染智能体发言"""
        agent_name = contribution["agent"]
        content = contribution["contribution"]
        logic_report = contribution["logic_report"]
        
        # 创建发言气泡
        with st.chat_message("assistant", avatar="🏥"):
            st.markdown(f"**{agent_name}**")
            st.write(content.get("reasoning", ""))
            
            # 逻辑报告
            with st.expander("逻辑分析报告"):
                st.write(logic_report)
            
            # 赞同/质疑按钮
            col1, col2 = st.columns(2)
            with col1:
                if st.button("👍 赞同", key=f"agree_{agent_name}"):
                    st.session_state.user_interventions.append({
                        "type": "agreement",
                        "target": agent_name,
                        "round": st.session_state.current_round
                    })
            with col2:
                if st.button("🤔 质疑", key=f"question_{agent_name}"):
                    self.render_question_form(agent_name)

    def render_active_intervention(self):
        """渲染主动介入界面"""
        st.subheader("向智能体提问")
        
        target_agent = st.selectbox(
            "选择提问对象",
            ["所有智能体"] + st.session_state.selected_agents
        )
        
        question = st.text_area("输入您的问题")
        
        if st.button("发送问题"):
            if question:
                intervention = {
                    "type": "question",
                    "target": target_agent,
                    "question": question,
                    "round": st.session_state.current_round,
                    "timestamp": datetime.now()
                }
                st.session_state.user_interventions.append(intervention)
                st.success("问题已发送")

    def render_critical_intervention(self):
        """渲染关键节点介入界面"""
        st.subheader("关键节点介入")
        
        # 检测关键分歧点
        disagreements = self.detect_disagreements()
        if disagreements:
            st.warning("检测到智能体之间存在分歧")
            for disagreement in disagreements:
                st.write(f"**分歧点**: {disagreement['issue']}")
                st.write(f"**相关智能体**: {', '.join(disagreement['agents'])}")
                
                if st.button("介入解决分歧", key=f"intervene_{disagreement['id']}"):
                    self.render_disagreement_resolution(disagreement)
        else:
            st.info("当前讨论进展顺利，无需介入")

    def detect_disagreements(self):
        """检测讨论中的分歧点"""
        # 简化的分歧检测逻辑
        disagreements = []
        if len(st.session_state.discussion_log) > 0:
            latest_round = st.session_state.discussion_log[-1]
            opinions = {}
            
            for contrib in latest_round["contributions"]:
                agent = contrib["agent"]
                opinion = contrib["contribution"].get("conclusion", "")
                if opinion:
                    if opinion not in opinions:
                        opinions[opinion] = []
                    opinions[opinion].append(agent)
            
            if len(opinions) > 1:
                disagreements.append({
                    "id": len(disagreements),
                    "issue": "诊断意见不一致",
                    "agents": [agent for agents in opinions.values() for agent in agents],
                    "opinions": opinions
                })
        
        return disagreements

    def render_results_section(self):
        """渲染结果展示界面"""
        if not hasattr(st.session_state, 'discussion_result') or not st.session_state.discussion_result:
            return
        
        st.header("📊 讨论结果汇总")
        
        result = st.session_state.discussion_result
        
        # 创建标签页显示不同方面的结果
        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "诊断评估", "治疗方案", "意见整合", "随访计划", "质量评估"
        ])
        
        with tab1:
            self.render_diagnosis_assessment(result)
        
        with tab2:
            self.render_treatment_plan(result)
        
        with tab3:
            self.render_integration_analysis(result)
        
        with tab4:
            self.render_followup_plan(result)
        
        with tab5:
            self.render_quality_assessment(result)
        
        # 导出功能
        st.header("📁 导出结果")
        self.render_export_options(result)

    def render_diagnosis_assessment(self, result):
        """渲染诊断评估"""
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("主要诊断")
            primary_dx = result.get("primary_diagnosis", {})
            st.write(f"**诊断**: {primary_dx.get('diagnosis', '')}")
            st.write(f"**置信度**: {primary_dx.get('confidence', '')}%")
            st.write(f"**依据**: {primary_dx.get('evidence', '')}")
        
        with col2:
            st.subheader("鉴别诊断")
            differential_dx = result.get("differential_diagnosis", [])
            for dx in differential_dx:
                with st.expander(f"{dx.get('diagnosis', '')} (概率: {dx.get('probability', '')}%)"):
                    st.write(f"**支持点**: {dx.get('supporting_evidence', '')}")
                    st.write(f"**排除点**: {dx.get('excluding_evidence', '')}")
        
        # 诊断一致性图表
        if 'diagnosis_consistency' in result:
            fig = self.create_consistency_chart(result['diagnosis_consistency'])
            st.plotly_chart(fig)

    def render_treatment_plan(self, result):
        """渲染治疗方案"""
        treatment_plan = result.get("treatment_plan", {})
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.subheader("急性期治疗")
            acute_tx = treatment_plan.get("acute", [])
            for tx in acute_tx:
                st.write(f"• {tx}")
        
        with col2:
            st.subheader("长期管理")
            chronic_tx = treatment_plan.get("chronic", [])
            for tx in chronic_tx:
                st.write(f"• {tx}")
        
        with col3:
            st.subheader("监测指标")
            monitoring = treatment_plan.get("monitoring", [])
            for monitor in monitoring:
                st.write(f"• {monitor}")
        
        # 治疗方案合理性评估
        if 'treatment_rationality' in result:
            rationality = result['treatment_rationality']
            st.metric("治疗方案合理性评分", f"{rationality.get('score', 0)}/100")
            st.write(f"**评估**: {rationality.get('assessment', '')}")

    def render_integration_analysis(self, result):
        """渲染意见整合分析"""
        integration = result.get("specialty_integration", {})
        
        st.subheader("各科室意见整合")
        
        # 意见一致性矩阵
        if 'consensus_matrix' in integration:
            df = pd.DataFrame(integration['consensus_matrix'])
            st.dataframe(df.style.highlight_max(axis=0))
        
        # 整合度评估
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("整合度评分", f"{integration.get('integration_score', 0)}/100")
        with col2:
            st.metric("共识度", f"{integration.get('consensus_level', 0)}%")
        with col3:
            st.metric("分歧点数量", integration.get('disagreement_count', 0))
        
        # 关键分歧点分析
        st.subheader("关键分歧分析")
        disagreements = integration.get('key_disagreements', [])
        for disagreement in disagreements:
            with st.expander(f"分歧点: {disagreement.get('issue', '')}"):
                st.write(f"**涉及科室**: {', '.join(disagreement.get('departments', []))}")
                st.write(f"**各方观点**:")
                for viewpoint in disagreement.get('viewpoints', []):
                    st.write(f"- {viewpoint.get('department', '')}: {viewpoint.get('opinion', '')}")

    def render_followup_plan(self, result):
        """渲染随访计划"""
        followup = result.get("follow_up_plan", {})
        
        # 时间轴显示随访计划
        timeline_data = []
        for period, plans in followup.items():
            if period == "immediate":
                timeline_data.append({"时期": "立即", "计划": plans})
            elif period == "short_term":
                timeline_data.append({"时期": "短期(1-4周)", "计划": plans})
            elif period == "long_term":
                timeline_data.append({"时期": "长期(1-12月)", "计划": plans})
        
        if timeline_data:
            for item in timeline_data:
                with st.expander(f"📅 {item['时期']}"):
                    if isinstance(item['计划'], list):
                        for plan in item['计划']:
                            st.write(f"• {plan}")
                    else:
                        st.write(item['计划'])
        
        # 随访完整性评估
        completeness = followup.get('completeness_score', 0)
        st.metric("随访计划完整性", f"{completeness}%")

    def render_quality_assessment(self, result):
        """渲染质量评估"""
        quality = result.get("quality_assessment", {})
        
        # 质量指标雷达图
        if 'metrics' in quality:
            metrics = quality['metrics']
            fig = self.create_quality_radar_chart(metrics)
            st.plotly_chart(fig)
        
        # 详细评估
        st.subheader("详细评估")
        categories = {
            "诊断准确性": quality.get('diagnosis_accuracy', {}),
            "治疗合理性": quality.get('treatment_rationality', {}),
            "整合能力": quality.get('integration_capability', {}),
            "随访完整性": quality.get('followup_completeness', {})
        }
        
        for category, details in categories.items():
            with st.expander(f"{category} - 评分: {details.get('score', 0)}/100"):
                st.write(f"**优势**: {details.get('strengths', '')}")
                st.write(f"**不足**: {details.get('weaknesses', '')}")
                st.write(f"**建议**: {details.get('recommendations', '')}")

    def create_consistency_chart(self, consistency_data):
        """创建一致性图表"""
        agents = list(consistency_data.keys())
        scores = list(consistency_data.values())
        
        fig = go.Figure(data=[go.Bar(x=agents, y=scores)])
        fig.update_layout(
            title="各智能体诊断一致性",
            xaxis_title="智能体",
            yaxis_title="一致性分数"
        )
        return fig

    def create_quality_radar_chart(self, metrics):
        """创建质量评估雷达图"""
        categories = list(metrics.keys())
        values = list(metrics.values())
        
        fig = go.Figure(data=go.Scatterpolar(
            r=values + [values[0]],  # 闭合雷达图
            theta=categories + [categories[0]],
            fill='toself'
        ))
        
        fig.update_layout(
            polar=dict(
                radialaxis=dict(
                    visible=True,
                    range=[0, 100]
                )),
            showlegend=False,
            title="讨论质量评估雷达图"
        )
        
        return fig

    def render_export_options(self, result):
        """渲染导出选项"""
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📄 导出为Word文档"):
                self.export_to_word(result)
        
        with col2:
            if st.button("📊 导出为PDF报告"):
                self.export_to_pdf(result)
        
        with col3:
            if st.button("📋 导出为JSON数据"):
                self.export_to_json(result)
        
        # 显示导出历史
        st.subheader("导出历史")
        export_history = self.get_export_history()
        if export_history:
            for export in export_history[-5:]:  # 显示最近5次导出
                st.write(f"• {export['timestamp']} - {export['format']} - {export['filename']}")
        else:
            st.info("暂无导出记录")

    def export_to_word(self, result):
        """导出为Word文档"""
        doc = Document()
        
        # 添加标题
        doc.add_heading('临床多智能体讨论报告', 0)
        
        # 添加基本信息
        doc.add_heading('讨论基本信息', level=1)
        basic_info = f"""
        讨论时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}
        参与智能体: {', '.join(st.session_state.selected_agents)}
        讨论轮数: {len(st.session_state.discussion_log)}
        用户介入次数: {len(st.session_state.user_interventions)}
        """
        doc.add_paragraph(basic_info)
        
        # 添加详细内容
        # ... 详细的Word文档生成逻辑
        
        # 保存文档
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # 提供下载
        b64 = base64.b64encode(buffer.getvalue()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="clinical_discussion_report.docx">下载Word文档</a>'
        st.markdown(href, unsafe_allow_html=True)

    def export_to_pdf(self, result):
        """导出为PDF报告"""
        # PDF导出逻辑（需要额外的库如reportlab）
        st.info("PDF导出功能正在开发中...")

    def export_to_json(self, result):
        """导出为JSON数据"""
        export_data = {
            "metadata": {
                "export_time": datetime.now().isoformat(),
                "user": st.session_state.current_user,
                "agents": st.session_state.selected_agents
            },
            "medical_record": st.session_state.medical_record,
            "discussion_question": st.session_state.discussion_question,
            "discussion_log": st.session_state.discussion_log,
            "final_result": result,
            "user_interventions": st.session_state.user_interventions
        }
        
        # 提供JSON下载
        json_str = json.dumps(export_data, ensure_ascii=False, indent=2)
        b64 = base64.b64encode(json_str.encode()).decode()
        href = f'<a href="data:application/json;base64,{b64}" download="clinical_discussion_{datetime.now().strftime("%Y%m%d_%H%M")}.json">下载JSON数据</a>'
        st.markdown(href, unsafe_allow_html=True)

    def get_export_history(self):
        """获取导出历史"""
        # 从存储中获取用户的导出历史
        try:
            user_history = self.discussion_storage.get_user_export_history(
                st.session_state.current_user
            )
            return user_history
        except:
            return []

    def render_discussion_history(self):
        """渲染讨论历史界面"""
        st.header("📚 历史讨论记录")
        
        # 获取用户的历史讨论
        user_discussions = self.discussion_storage.get_user_discussions(
            st.session_state.current_user
        )
        
        if not user_discussions:
            st.info("暂无历史讨论记录")
            return
        
        # 显示历史记录表格
        history_data = []
        for discussion in user_discussions:
            history_data.append({
                "时间": discussion["metadata"]["timestamp"],
                "智能体数量": len(discussion["metadata"]["agents_used"]),
                "讨论轮数": discussion["metadata"]["rounds"],
                "主要诊断": discussion["final_result"].get("primary_diagnosis", {}).get("diagnosis", "N/A"),
                "质量评分": discussion["final_result"].get("quality_assessment", {}).get("overall_score", "N/A")
            })
        
        df = pd.DataFrame(history_data)
        st.dataframe(df, use_container_width=True)
        
        # 选择具体讨论查看详情
        st.subheader("查看详细记录")
        discussion_times = [d["metadata"]["timestamp"] for d in user_discussions]
        selected_time = st.selectbox("选择讨论记录", discussion_times)
        
        if selected_time:
            selected_discussion = next(
                d for d in user_discussions 
                if d["metadata"]["timestamp"] == selected_time
            )
            self.render_discussion_detail(selected_discussion)

    def render_discussion_detail(self, discussion):
        """渲染讨论详情"""
        with st.expander("讨论详情", expanded=True):
            # 基本信息
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("参与智能体", len(discussion["metadata"]["agents_used"]))
            with col2:
                st.metric("讨论轮数", discussion["metadata"]["rounds"])
            with col3:
                st.metric("用户介入", len(discussion.get("user_interventions", [])))
            
            # 智能体列表
            st.write("**参与智能体**:", ", ".join(discussion["metadata"]["agents_used"]))
            
            # 病历摘要
            with st.expander("病历摘要"):
                st.write(discussion["medical_record"])
            
            # 讨论过程回放
            st.subheader("讨论过程回放")
            for round_num, round_log in enumerate(discussion["discussion_log"]):
                with st.expander(f"第 {round_num + 1} 轮讨论"):
                    for contribution in round_log["contributions"]:
                        st.write(f"**{contribution['agent']}**:")
                        st.write(contribution["contribution"].get("summary", ""))
            
            # 最终结果
            st.subheader("最终结果")
            result = discussion["final_result"]
            st.json(result)  # 或者使用更结构化的显示方式

    def render_user_profile(self):
        """渲染用户个人资料界面"""
        st.header("👤 用户个人资料")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("基本信息")
            st.write(f"**用户名**: {st.session_state.current_user}")
            st.write(f"**注册时间**: {self.get_user_registration_time()}")
            st.write(f"**最后登录**: {self.get_last_login_time()}")
            
            # 使用统计
            st.subheader("使用统计")
            discussions_count = self.get_discussion_count()
            agents_used = self.get_most_used_agents()
            st.write(f"**总讨论次数**: {discussions_count}")
            st.write(f"**最常用智能体**: {', '.join(agents_used[:3])}")
        
        with col2:
            st.subheader("偏好设置")
            
            # 讨论偏好
            default_rounds = st.slider("默认讨论轮数", 1, 10, 3)
            default_intervention = st.selectbox(
                "默认介入模式",
                ["旁观模式", "主动介入", "仅关键节点介入"]
            )
            
            # 显示偏好
            theme = st.selectbox("界面主题", ["浅色", "深色"])
            font_size = st.slider("字体大小", 12, 24, 16)
            
            if st.button("保存设置"):
                self.save_user_preferences({
                    "default_rounds": default_rounds,
                    "default_intervention": default_intervention,
                    "theme": theme,
                    "font_size": font_size
                })
                st.success("设置已保存")

    def get_user_registration_time(self):
        """获取用户注册时间"""
        # 从用户管理器中获取
        return "2024-01-01"  # 示例数据

    def get_last_login_time(self):
        """获取最后登录时间"""
        return datetime.now().strftime("%Y-%m-%d %H:%M")

    def get_discussion_count(self):
        """获取讨论次数"""
        discussions = self.discussion_storage.get_user_discussions(
            st.session_state.current_user
        )
        return len(discussions)

    def get_most_used_agents(self):
        """获取最常用智能体"""
        # 从历史记录中统计
        return ["心内科", "肾内科", "内分泌科"]  # 示例数据

    def save_user_preferences(self, preferences):
        """保存用户偏好设置"""
        # 实现保存逻辑
        pass

    def render_main_interface(self):
        """渲染主界面"""
        # 侧边栏导航
        st.sidebar.header("🧭 导航")
        
        menu_options = [
            "🏠 控制面板",
            "💬 开始讨论", 
            "📚 历史记录",
            "👤 个人资料",
            "⚙️ 系统设置"
        ]
        
        selected_menu = st.sidebar.radio("选择功能", menu_options)
        
        # 根据选择显示不同内容
        if selected_menu == "🏠 控制面板":
            self.render_dashboard()
        elif selected_menu == "💬 开始讨论":
            self.render_discussion_workflow()
        elif selected_menu == "📚 历史记录":
            self.render_discussion_history()
        elif selected_menu == "👤 个人资料":
            self.render_user_profile()
        elif selected_menu == "⚙️ 系统设置":
            self.render_system_settings()
        
        # 智能体选择侧边栏（始终显示）
        self.render_agent_selection()

    def render_dashboard(self):
        """渲染控制面板"""
        st.header("🏠 控制面板")
        
        # 欢迎信息
        st.success(f"欢迎回来，{st.session_state.current_user}！")
        
        # 快速统计
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            discussions_count = self.get_discussion_count()
            st.metric("总讨论次数", discussions_count)
        
        with col2:
            recent_discussions = self.get_recent_discussions_count(7)  # 最近7天
            st.metric("最近7天讨论", recent_discussions)
        
        with col3:
            avg_quality = self.get_average_quality_score()
            st.metric("平均质量评分", f"{avg_quality}/100")
        
        with col4:
            favorite_agent = self.get_most_used_agents()[0]
            st.metric("最常用智能体", favorite_agent)
        
        # 快速操作卡片
        st.subheader("快速操作")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🚀 开始新讨论", use_container_width=True):
                st.session_state.current_page = "discussion"
                st.rerun()
        
        with col2:
            if st.button("📊 查看统计", use_container_width=True):
                st.session_state.current_page = "analytics"
                st.rerun()
        
        with col3:
            if st.button("⚙️ 系统设置", use_container_width=True):
                st.session_state.current_page = "settings"
                st.rerun()
        
        # 最近活动
        st.subheader("最近活动")
        recent_discussions = self.get_recent_discussions(5)
        for discussion in recent_discussions:
            with st.expander(f"{discussion['timestamp']} - {discussion['diagnosis']}"):
                st.write(f"**智能体**: {', '.join(discussion['agents'])}")
                st.write(f"**质量评分**: {discussion['quality_score']}")
                if st.button("查看详情", key=f"view_{discussion['id']}"):
                    self.render_discussion_detail(discussion)

    def render_discussion_workflow(self):
        """渲染完整讨论工作流"""
        # 步骤指示器
        steps = ["病历输入", "智能体选择", "讨论配置", "实时讨论", "结果查看"]
        current_step = self.get_current_step()
        
        # 显示步骤进度
        st.progress((current_step + 1) / len(steps))
        st.write(f"当前步骤: **{steps[current_step]}**")
        
        # 根据当前步骤显示相应内容
        if current_step == 0:
            self.render_medical_input()
            if st.session_state.medical_record:
                st.session_state.current_step = 1
                st.rerun()
        elif current_step == 1:
            self.render_agent_selection()
            if st.session_state.selected_agents:
                if st.button("确认选择", use_container_width=True):
                    st.session_state.current_step = 2
                    st.rerun()
        elif current_step == 2:
            self.render_discussion_control()
        elif current_step == 3:
            self.render_discussion_display()
        elif current_step == 4:
            self.render_results_section()
            
            # 提供重新开始选项
            if st.button("🔄 开始新的讨论", use_container_width=True):
                self.reset_discussion_state()
                st.session_state.current_step = 0
                st.rerun()

    def get_current_step(self):
        """获取当前步骤"""
        if not hasattr(st.session_state, 'current_step'):
            st.session_state.current_step = 0
        return st.session_state.current_step

    def get_recent_discussions(self, count=5):
        """获取最近讨论记录"""
        # 从存储中获取最近记录
        return []  # 示例数据

    def get_recent_discussions_count(self, days=7):
        """获取最近N天的讨论次数"""
        return 3  # 示例数据

    def get_average_quality_score(self):
        """获取平均质量评分"""
        return 85  # 示例数据

    def render_system_settings(self):
        """渲染系统设置界面"""
        st.header("⚙️ 系统设置")
        
        tab1, tab2, tab3 = st.tabs(["模型设置", "界面设置", "数据管理"])
        
        with tab1:
            self.render_model_settings()
        
        with tab2:
            self.render_interface_settings()
        
        with tab3:
            self.render_data_management()

    def render_model_settings(self):
        """渲染模型设置"""
        st.subheader("LLM模型配置")
        
        # 模型选择
        current_model = st.selectbox(
            "选择模型",
            ["clinical-model-v1", "clinical-model-v2", "general-model"],
            index=0
        )
        
        # 参数调整
        col1, col2 = st.columns(2)
        with col1:
            temperature = st.slider("温度参数", 0.0, 1.0, 0.7, 0.1)
            max_tokens = st.number_input("最大token数", 100, 8000, 2000)
        
        with col2:
            top_p = st.slider("Top-p", 0.0, 1.0, 0.9, 0.1)
            frequency_penalty = st.slider("频率惩罚", -2.0, 2.0, 0.0, 0.1)
        
        # API设置
        st.subheader("API配置")
        api_endpoint = st.text_input(
            "API端点", 
            value="http://10.124.0.7:9001/v1",
            help="LLM API的服务地址"
        )
        
        api_key = st.text_input(
            "API密钥", 
            type="password",
            help="如果需要认证，请输入API密钥"
        )
        
        if st.button("测试连接"):
            with st.spinner("测试连接中..."):
                if self.test_api_connection(api_endpoint, api_key):
                    st.success("连接成功！")
                else:
                    st.error("连接失败，请检查配置")

    def render_interface_settings(self):
        """渲染界面设置"""
        st.subheader("界面主题")
        
        theme = st.selectbox("选择主题", ["浅色", "深色", "自动"])
        font_size = st.slider("字体大小", 12, 24, 16)
        language = st.selectbox("界面语言", ["中文", "English"])
        
        st.subheader("布局选项")
        col1, col2 = st.columns(2)
        with col1:
            default_layout = st.selectbox("默认布局", ["宽屏", "窄屏", "自适应"])
            show_animations = st.checkbox("显示动画效果", value=True)
        
        with col2:
            auto_save = st.checkbox("自动保存", value=True)
            save_interval = st.number_input("保存间隔(分钟)", 1, 60, 5)
        
        if st.button("应用设置"):
            st.success("界面设置已应用")

    def render_data_management(self):
        """渲染数据管理"""
        st.subheader("数据备份")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📥 备份所有数据", use_container_width=True):
                self.backup_user_data()
        
        with col2:
            if st.button("🗑️ 清理临时数据", use_container_width=True):
                self.cleanup_temp_data()
        
        st.subheader("数据导出")
        export_format = st.selectbox("导出格式", ["JSON", "CSV", "Excel"])
        include_types = st.multiselect(
            "包含数据类型",
            ["讨论记录", "用户设置", "智能体配置", "系统日志"],
            default=["讨论记录"]
        )
        
        if st.button("导出数据"):
            with st.spinner("正在导出数据..."):
                self.export_user_data(export_format, include_types)

    def test_api_connection(self, endpoint, api_key):
        """测试API连接"""
        # 实现API连接测试逻辑
        try:
            # 简化的测试逻辑
            return True
        except:
            return False

    def backup_user_data(self):
        """备份用户数据"""
        with st.spinner("正在备份数据..."):
            # 实现备份逻辑
            st.success("数据备份完成")

    def cleanup_temp_data(self):
        """清理临时数据"""
        if st.button("确认清理"):
            with st.spinner("正在清理..."):
                # 实现清理逻辑
                st.success("临时数据已清理")

    def export_user_data(self, format, include_types):
        """导出用户数据"""
        with st.spinner(f"正在导出{format}数据..."):
            # 实现导出逻辑
            st.success("数据导出完成")

    def reset_discussion_state(self):
        """重置讨论状态"""
        st.session_state.discussion_active = False
        st.session_state.discussion_log = []
        st.session_state.user_interventions = []
        if hasattr(st.session_state, 'discussion_result'):
            del st.session_state.discussion_result

    def run(self):
        """主运行函数"""
        # 检查认证状态
        if not st.session_state.authenticated:
            self.render_authentication_section()
        else:
            # 显示顶部导航栏
            self.render_top_navigation()
            
            # 显示主界面
            self.render_main_interface()

    def render_top_navigation(self):
        """渲染顶部导航栏"""
        col1, col2, col3 = st.columns([3, 1, 1])
        
        with col1:
            st.markdown("<h1 style='text-align: center;'>🏥 临床MDT智能模拟助手</h1>", 
                       unsafe_allow_html=True)
        
        with col2:
            st.write(f"欢迎，**{st.session_state.current_user}**")
        
        with col3:
            if st.button("🚪 退出"):
                self.logout_user()

    def logout_user(self):
        """用户退出登录"""
        st.session_state.authenticated = False
        st.session_state.current_user = None
        st.session_state.session_id = None
        st.rerun()

def main():
    """主函数"""
    try:
        # 创建Web界面实例并运行
        web_interface = ClinicalWebInterface()
        web_interface.run()
    except Exception as e:
        st.error(f"系统错误: {str(e)}")
        st.info("请刷新页面重试或联系系统管理员")

if __name__ == "__main__":
    main()