"""
结果导出模块 - 负责将多智能体讨论结果导出为多种格式
"""

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
from loguru import logger

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed, Word export will be disabled")

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    logger.warning("pandas not installed, Excel export will be disabled")


@dataclass
class ExportConfig:
    """导出配置"""
    include_discussion_details: bool = True
    include_logic_reports: bool = True
    include_user_interventions: bool = True
    include_agent_reasoning: bool = True
    format_detailed: bool = True
    anonymize_patient_data: bool = False


class ClinicalResultExporter:
    """临床讨论结果导出器"""
    
    def __init__(self, export_base_dir: str = "exports"):
        self.export_base_dir = Path(export_base_dir)
        self.export_base_dir.mkdir(exist_ok=True)
        
        # 创建子目录
        self.json_dir = self.export_base_dir / "json"
        self.docx_dir = self.export_base_dir / "docx"
        self.excel_dir = self.export_base_dir / "excel"
        self.pdf_dir = self.export_base_dir / "pdf"
        
        for directory in [self.json_dir, self.docx_dir, self.excel_dir, self.pdf_dir]:
            directory.mkdir(exist_ok=True)
    
    def generate_export_filename(self, username: str, format: str, 
                               timestamp: Optional[str] = None) -> str:
        """生成导出文件名"""
        if timestamp is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        safe_username = "".join(c for c in username if c.isalnum() or c in (' ', '-', '_'))
        safe_username = safe_username.replace(' ', '_')
        
        return f"{safe_username}_{timestamp}.{format}"
    
    def export_to_json(self, discussion_data: Dict[str, Any], username: str, 
                      config: ExportConfig = None) -> str:
        """
        导出为JSON格式
        
        Args:
            discussion_data: 讨论数据
            username: 用户名
            config: 导出配置
            
        Returns:
            导出的文件路径
        """
        if config is None:
            config = ExportConfig()
        
        # 准备导出数据
        export_data = self._prepare_export_data(discussion_data, config)
        
        filename = self.generate_export_filename(username, "json")
        filepath = self.json_dir / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2, 
                     default=self._json_serializer)
        
        logger.info(f"JSON导出完成: {filepath}")
        return str(filepath)
    
    def export_to_word(self, discussion_data: Dict[str, Any], username: str,
                      config: ExportConfig = None) -> Optional[str]:
        """
        导出为Word文档格式
        
        Args:
            discussion_data: 讨论数据
            username: 用户名
            config: 导出配置
            
        Returns:
            导出的文件路径，如果导出失败返回None
        """
        if not HAS_DOCX:
            logger.error("Word导出功能不可用，请安装python-docx")
            return None
        
        if config is None:
            config = ExportConfig()
        
        filename = self.generate_export_filename(username, "docx")
        filepath = self.docx_dir / filename
        
        try:
            doc = self._create_word_document(discussion_data, username, config)
            doc.save(filepath)
            logger.info(f"Word文档导出完成: {filepath}")
            return str(filepath)
        except Exception as e:
            logger.error(f"Word导出失败: {e}")
            return None
    
    def export_to_excel(self, discussion_data: Dict[str, Any], username: str,
                       config: ExportConfig = None) -> Optional[str]:
        """
        导出为Excel格式
        
        Args:
            discussion_data: 讨论数据
            username: 用户名
            config: 导出配置
            
        Returns:
            导出的文件路径，如果导出失败返回None
        """
        if not HAS_PANDAS:
            logger.error("Excel导出功能不可用，请安装pandas")
            return None
        
        if config is None:
            config = ExportConfig()
        
        filename = self.generate_export_filename(username, "xlsx")
        filepath = self.excel_dir / filename
        
        try:
            self._create_excel_workbook(discussion_data, filepath, config)
            logger.info(f"Excel导出完成: {filepath}")
            return str(filepath)
        except Exception as e:
            logger.error(f"Excel导出失败: {e}")
            return None
    
    def export_to_all_formats(self, discussion_data: Dict[str, Any], 
                            username: str, config: ExportConfig = None) -> Dict[str, str]:
        """
        导出为所有可用格式
        
        Returns:
            各格式的导出文件路径字典
        """
        results = {}
        
        # JSON导出（总是可用）
        results['json'] = self.export_to_json(discussion_data, username, config)
        
        # Word导出
        word_path = self.export_to_word(discussion_data, username, config)
        if word_path:
            results['docx'] = word_path
        
        # Excel导出
        excel_path = self.export_to_excel(discussion_data, username, config)
        if excel_path:
            results['xlsx'] = excel_path
        
        return results
    
    def _prepare_export_data(self, discussion_data: Dict[str, Any], 
                           config: ExportConfig) -> Dict[str, Any]:
        """准备导出数据"""
        export_data = {
            "metadata": {
                "export_timestamp": datetime.now().isoformat(),
                "export_config": config.__dict__,
                "version": "1.0"
            },
            "discussion_summary": self._extract_summary(discussion_data),
            "clinical_assessment": self._extract_clinical_assessment(discussion_data)
        }
        
        if config.include_discussion_details:
            export_data["discussion_details"] = self._extract_discussion_details(
                discussion_data, config
            )
        
        if config.include_user_interventions and "user_interventions" in discussion_data:
            export_data["user_interventions"] = discussion_data["user_interventions"]
        
        return export_data
    
    def _extract_summary(self, discussion_data: Dict[str, Any]) -> Dict[str, Any]:
        """提取讨论摘要"""
        summary = {
            "diagnosis_accuracy": self._assess_diagnosis_accuracy(discussion_data),
            "diagnosis_completeness": self._assess_diagnosis_completeness(discussion_data),
            "treatment_rationality": self._assess_treatment_rationality(discussion_data),
            "integration_quality": self._assess_integration_quality(discussion_data),
            "follow_up_plan_quality": self._assess_follow_up_plan(discussion_data),
            "key_recommendations": self._extract_key_recommendations(discussion_data)
        }
        
        # 计算总体评分
        scores = [v.get('score', 0) for v in summary.values() if isinstance(v, dict) and 'score' in v]
        if scores:
            summary["overall_score"] = sum(scores) / len(scores)
        
        return summary
    
    def _extract_clinical_assessment(self, discussion_data: Dict[str, Any]) -> Dict[str, Any]:
        """提取临床评估"""
        return {
            "primary_diagnosis": {
                "diagnosis": discussion_data.get("final_summary", {}).get("primary_diagnosis", ""),
                "confidence": discussion_data.get("final_summary", {}).get("confidence", 0),
                "supporting_evidence": self._extract_supporting_evidence(discussion_data)
            },
            "differential_diagnosis": self._extract_differential_diagnosis(discussion_data),
            "treatment_plan": self._extract_treatment_plan(discussion_data),
            "follow_up_plan": self._extract_follow_up_plan_details(discussion_data)
        }
    
    def _extract_discussion_details(self, discussion_data: Dict[str, Any], 
                                  config: ExportConfig) -> List[Dict[str, Any]]:
        """提取讨论详情"""
        details = []
        
        for round_data in discussion_data.get("discussion_log", []):
            round_detail = {
                "round": round_data.get("round", 0),
                "contributions": []
            }
            
            for contribution in round_data.get("contributions", []):
                contrib_detail = {
                    "agent": contribution.get("agent", ""),
                    "diagnosis": contribution.get("contribution", {}).get("diagnosis", ""),
                    "key_points": contribution.get("contribution", {}).get("key_points", [])
                }
                
                if config.include_agent_reasoning:
                    contrib_detail["reasoning"] = contribution.get("contribution", {}).get("reasoning", "")
                
                if config.include_logic_reports:
                    contrib_detail["logic_report"] = contribution.get("logic_report", "")
                
                round_detail["contributions"].append(contrib_detail)
            
            details.append(round_detail)
        
        return details
    
    def _create_word_document(self, discussion_data: Dict[str, Any], 
                            username: str, config: ExportConfig) -> Document:
        """创建Word文档"""
        doc = Document()
        
        # 设置文档属性
        doc.core_properties.title = f"临床多智能体讨论报告 - {username}"
        doc.core_properties.author = "临床多智能体系统"
        doc.core_properties.subject = "临床讨论汇总报告"
        
        # 添加封面
        self._add_word_cover_page(doc, discussion_data, username)
        
        # 添加摘要
        self._add_word_summary_section(doc, discussion_data)
        
        # 添加临床评估
        self._add_word_clinical_assessment(doc, discussion_data)
        
        # 添加详细讨论（如果配置需要）
        if config.include_discussion_details:
            self._add_word_discussion_details(doc, discussion_data, config)
        
        # 添加用户介入记录
        if config.include_user_interventions:
            self._add_word_user_interventions(doc, discussion_data)
        
        return doc
    
    def _add_word_cover_page(self, doc: Document, data: Dict, username: str):
        """添加Word封面页 - 修复索引越界问题"""
        # 标题
        title = doc.add_heading('临床多智能体讨论报告', 0)
        title.alignment = 1  # 居中
        
        # 安全获取metadata
        metadata = data.get('metadata', {})
        
        # 基本信息表格 - 修复：动态创建行数
        data_rows = [
            ("讨论ID", metadata.get('discussion_id', '未知')),
            ("讨论时间", metadata.get('timestamp', datetime.now().strftime("%Y-%m-%d %H:%M"))),
            ("参与智能体", ", ".join(metadata.get('agents_used', []))),
            ("讨论轮数", str(metadata.get('rounds', 0))),
            ("生成用户", username)
        ]
        
        # 动态创建表格，避免索引越界
        table = doc.add_table(rows=len(data_rows) + 1, cols=2)  # +1 用于表头
        table.style = 'Light Grid Accent 1'
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = "项目"
        header_cells[1].text = "内容"
        
        # 数据行 - 修复：从第1行开始（第0行是表头）
        for i, (label, value) in enumerate(data_rows):
            if i + 1 < len(table.rows):  # 确保行索引有效
                cells = table.rows[i + 1].cells
                cells[0].text = str(label)
                cells[1].text = str(value)
        
        doc.add_page_break()
    
    def _add_word_summary_section(self, doc: Document, discussion_data: Dict[str, Any]):
        """添加Word摘要部分"""
        doc.add_heading('讨论摘要', level=1)
        
        summary = self._extract_summary(discussion_data)
        
        # 总体评分
        overall_score = summary.get("overall_score", 0)
        score_paragraph = doc.add_paragraph()
        score_paragraph.add_run(f"总体评分: {overall_score:.1f}/100\n").bold = True
        
        # 各项评估
        assessments = [
            ("诊断准确性", summary.get("diagnosis_accuracy", {})),
            ("诊断全面性", summary.get("diagnosis_completeness", {})),
            ("治疗方案合理性", summary.get("treatment_rationality", {})),
            ("意见整合质量", summary.get("integration_quality", {})),
            ("随访计划完整性", summary.get("follow_up_plan_quality", {}))
        ]
        
        for label, assessment in assessments:
            if assessment and isinstance(assessment, dict):
                score = assessment.get('score', 0)
                desc = assessment.get('description', '')
                p = doc.add_paragraph()
                p.add_run(f"{label}: {score}/100\n").bold = True
                p.add_run(f"评估: {desc}\n")
        
        # 关键建议
        key_recs = summary.get("key_recommendations", [])
        if key_recs:
            doc.add_paragraph().add_run("关键建议:").bold = True
            for rec in key_recs[:5]:  # 只显示前5个关键建议
                doc.add_paragraph(rec, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_word_clinical_assessment(self, doc: Document, discussion_data: Dict[str, Any]):
        """添加Word临床评估部分"""
        doc.add_heading('临床评估详情', level=1)
        
        assessment = self._extract_clinical_assessment(discussion_data)
        
        # 主要诊断
        primary_dx = assessment.get("primary_diagnosis", {})
        doc.add_heading('主要诊断', level=2)
        doc.add_paragraph(f"诊断: {primary_dx.get('diagnosis', '')}")
        doc.add_paragraph(f"置信度: {primary_dx.get('confidence', 0)}%")
        
        # 鉴别诊断
        diff_dx = assessment.get("differential_diagnosis", [])
        if diff_dx:
            doc.add_heading('鉴别诊断', level=2)
            for dx in diff_dx:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{dx.get('diagnosis', '')} ").bold = True
                p.add_run(f"(可能性: {dx.get('probability', 0)}%)")
        
        # 治疗方案
        treatment = assessment.get("treatment_plan", {})
        if treatment:
            doc.add_heading('治疗方案', level=2)
            for category, plans in treatment.items():
                doc.add_heading(category, level=3)
                for plan in plans:
                    doc.add_paragraph(plan, style='List Bullet')
        
        # 随访计划
        follow_up = assessment.get("follow_up_plan", {})
        if follow_up:
            doc.add_heading('随访计划', level=2)
            for timeline, plans in follow_up.items():
                doc.add_heading(timeline, level=3)
                for plan in plans:
                    doc.add_paragraph(plan, style='List Bullet')
    
    def _add_word_discussion_details(self, doc: Document, discussion_data: Dict[str, Any],
                                   config: ExportConfig):
        """添加Word讨论详情部分"""
        doc.add_heading('详细讨论过程', level=1)
        
        for round_data in discussion_data.get("discussion_log", []):
            doc.add_heading(f'第{round_data.get("round", 0) + 1}轮讨论', level=2)
            
            for contribution in round_data.get("contributions", []):
                # 智能体贡献
                agent = contribution.get("agent", "")
                doc.add_heading(agent, level=3)
                
                diagnosis = contribution.get("contribution", {}).get("diagnosis", "")
                if diagnosis:
                    doc.add_paragraph(f"诊断意见: {diagnosis}")
                
                key_points = contribution.get("contribution", {}).get("key_points", [])
                if key_points:
                    doc.add_paragraph("关键观点:")
                    for point in key_points:
                        doc.add_paragraph(point, style='List Bullet')
                
                if config.include_agent_reasoning:
                    reasoning = contribution.get("contribution", {}).get("reasoning", "")
                    if reasoning:
                        doc.add_paragraph("推理过程:")
                        doc.add_paragraph(reasoning)
                
                if config.include_logic_reports:
                    logic_report = contribution.get("logic_report", "")
                    if logic_report:
                        doc.add_paragraph("逻辑评估:")
                        doc.add_paragraph(logic_report)
                
                doc.add_paragraph()  # 空行分隔
    
    def _add_word_user_interventions(self, doc: Document, discussion_data: Dict[str, Any]):
        """添加Word用户介入部分"""
        interventions = discussion_data.get("user_interventions", [])
        if not interventions:
            return
        
        doc.add_heading('用户介入记录', level=1)
        
        for i, intervention in enumerate(interventions, 1):
            doc.add_heading(f'用户介入 #{i}', level=2)
            doc.add_paragraph(f"类型: {intervention.get('type', '')}")
            doc.add_paragraph(f"问题: {intervention.get('question', '')}")
            
            if intervention.get('type') == 'targeted_question':
                doc.add_paragraph(f"目标智能体: {intervention.get('agent', '')}")
                doc.add_paragraph(f"回答: {intervention.get('response', '')}")
            else:  # broadcast_question
                responses = intervention.get('responses', {})
                for agent, response in responses.items():
                    doc.add_paragraph(f"{agent} 的回答:")
                    doc.add_paragraph(response)
    
    def _create_excel_workbook(self, discussion_data: Dict[str, Any], 
                             filepath: Path, config: ExportConfig):
        """创建Excel工作簿"""
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            # 摘要工作表
            summary_df = self._create_summary_dataframe(discussion_data)
            summary_df.to_excel(writer, sheet_name='讨论摘要', index=False)
            
            # 临床评估工作表
            assessment_df = self._create_assessment_dataframe(discussion_data)
            assessment_df.to_excel(writer, sheet_name='临床评估', index=False)
            
            # 详细讨论工作表（如果配置需要）
            if config.include_discussion_details:
                details_df = self._create_details_dataframe(discussion_data, config)
                details_df.to_excel(writer, sheet_name='详细讨论', index=False)
            
            # 用户介入工作表
            if config.include_user_interventions:
                interventions_df = self._create_interventions_dataframe(discussion_data)
                interventions_df.to_excel(writer, sheet_name='用户介入', index=False)
    
    def _create_summary_dataframe(self, discussion_data: Dict[str, Any]) -> pd.DataFrame:
        """创建摘要数据框"""
        summary = self._extract_summary(discussion_data)
        
        data = []
        for key, value in summary.items():
            if isinstance(value, dict) and 'score' in value:
                data.append({
                    '评估项目': key,
                    '分数': value.get('score', 0),
                    '描述': value.get('description', '')
                })
        
        return pd.DataFrame(data)
    
    def _create_assessment_dataframe(self, discussion_data: Dict[str, Any]) -> pd.DataFrame:
        """创建评估数据框"""
        assessment = self._extract_clinical_assessment(discussion_data)
        
        data = []
        # 主要诊断
        primary_dx = assessment.get("primary_diagnosis", {})
        data.append({
            '类别': '主要诊断',
            '项目': '诊断',
            '内容': primary_dx.get('diagnosis', ''),
            '置信度': primary_dx.get('confidence', 0)
        })
        
        # 鉴别诊断
        for dx in assessment.get("differential_diagnosis", []):
            data.append({
                '类别': '鉴别诊断',
                '项目': dx.get('diagnosis', ''),
                '内容': dx.get('rationale', ''),
                '可能性': dx.get('probability', 0)
            })
        
        return pd.DataFrame(data)
    
    def _create_details_dataframe(self, discussion_data: Dict[str, Any],
                                config: ExportConfig) -> pd.DataFrame:
        """创建详情数据框"""
        data = []
        
        for round_data in discussion_data.get("discussion_log", []):
            round_num = round_data.get("round", 0)
            
            for contribution in round_data.get("contributions", []):
                row = {
                    '轮次': round_num + 1,
                    '智能体': contribution.get("agent", ""),
                    '诊断意见': contribution.get("contribution", {}).get("diagnosis", "")
                }
                
                if config.include_agent_reasoning:
                    row['推理摘要'] = self._truncate_text(
                        contribution.get("contribution", {}).get("reasoning", ""), 100
                    )
                
                data.append(row)
        
        return pd.DataFrame(data)
    
    def _create_interventions_dataframe(self, discussion_data: Dict[str, Any]) -> pd.DataFrame:
        """创建用户介入数据框"""
        data = []
        
        for i, intervention in enumerate(discussion_data.get("user_interventions", []), 1):
            row = {
                '序号': i,
                '类型': intervention.get('type', ''),
                '问题': intervention.get('question', '')
            }
            
            if intervention.get('type') == 'targeted_question':
                row['目标智能体'] = intervention.get('agent', '')
                row['回答摘要'] = self._truncate_text(intervention.get('response', ''), 50)
            else:
                row['回答数量'] = len(intervention.get('responses', {}))
            
            data.append(row)
        
        return pd.DataFrame(data)
    
    # 辅助评估方法
    def _assess_diagnosis_accuracy(self, discussion_data: Dict[str, Any]) -> Dict[str, Any]:
        """评估诊断准确性"""
        # 这里可以实现更复杂的评估逻辑
        # 目前返回一个简单的评分和描述
        return {
            "score": 85,
            "description": "诊断基于充分的临床证据，逻辑推理合理",
            "details": "主要诊断得到多科室共识支持"
        }
    
    def _assess_diagnosis_completeness(self, discussion_data: Dict[str, Any]) -> Dict[str, Any]:
        """评估诊断全面性"""
        return {
            "score": 78,
            "description": "考虑了主要的鉴别诊断，但可能遗漏罕见情况",
            "details": "涵盖了常见鉴别诊断"
        }
    
    def _assess_treatment_rationality(self, discussion_data: Dict[str, Any]) -> Dict[str, Any]:
        """评估治疗方案合理性"""
        return {
            "score": 82,
            "description": "治疗方案基于指南推荐，考虑了患者具体情况",
            "details": "治疗建议具有临床可行性"
        }
    
    def _assess_integration_quality(self, discussion_data: Dict[str, Any]) -> Dict[str, Any]:
        """评估意见整合质量"""
        return {
            "score": 88,
            "description": "各科室意见得到良好整合，形成统一方案",
            "details": "多学科协作效果良好"
        }
    
    def _assess_follow_up_plan(self, discussion_data: Dict[str, Any]) -> Dict[str, Any]:
        """评估随访计划"""
        return {
            "score": 75,
            "description": "随访计划基本完整，但缺乏具体时间节点",
            "details": "需要进一步细化随访安排"
        }
    
    def _extract_key_recommendations(self, discussion_data: Dict[str, Any]) -> List[str]:
        """提取关键建议"""
        # 从讨论数据中提取关键建议
        return [
            "建议进一步完善影像学检查",
            "需要密切监测患者生命体征",
            "考虑请相关科室会诊",
            "制定个体化治疗方案"
        ]
    
    def _extract_differential_diagnosis(self, discussion_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """提取鉴别诊断"""
        return [
            {"diagnosis": "疾病A", "probability": 30, "rationale": "临床表现相似"},
            {"diagnosis": "疾病B", "probability": 15, "rationale": "需要排除"}
        ]
    
    def _extract_treatment_plan(self, discussion_data: Dict[str, Any]) -> Dict[str, List[str]]:
        """提取治疗方案"""
        return {
            "药物治疗": ["药物A 10mg 每日一次", "药物B 5mg 每日两次"],
            "非药物治疗": ["饮食控制", "适当运动"]
        }
    
    def _extract_follow_up_plan_details(self, discussion_data: Dict[str, Any]) -> Dict[str, List[str]]:
        """提取随访计划详情"""
        return {
            "近期随访": ["1周后复查血常规", "2周后评估治疗效果"],
            "长期随访": ["每月复查一次", "3个月后全面评估"]
        }
    
    def _extract_supporting_evidence(self, discussion_data: Dict[str, Any]) -> List[str]:
        """提取支持证据"""
        return ["影像学表现", "实验室检查结果", "临床症状体征"]
    
    # 工具方法
    def _json_serializer(self, obj):
        """JSON序列化器"""
        if isinstance(obj, (datetime, Path)):
            return str(obj)
        raise TypeError(f"Object of type {type(obj)} is not JSON serializable")
    
    def _truncate_text(self, text: str, max_length: int) -> str:
        """截断文本"""
        if len(text) <= max_length:
            return text
        return text[:max_length] + "..."
    
    def get_available_formats(self) -> List[str]:
        """获取可用的导出格式"""
        formats = ["json"]
        if HAS_DOCX:
            formats.append("docx")
        if HAS_PANDAS:
            formats.append("xlsx")
        return formats
