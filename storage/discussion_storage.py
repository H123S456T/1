import json
import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any
from loguru import logger
import pandas as pd
from docx import Document
from docx.shared import Inches
import html

class DiscussionStorage:
    """讨论记录存储和管理类"""
    
    def __init__(self, storage_base: str = "data/discussions"):
        self.storage_base = Path(storage_base)
        self.storage_base.mkdir(parents=True, exist_ok=True)
        
        # 创建必要的子目录
        (self.storage_base / "json").mkdir(exist_ok=True)
        (self.storage_base / "exports").mkdir(exist_ok=True)
        (self.storage_base / "backups").mkdir(exist_ok=True)
   
    def get_user_discussions(self, user_id: str, limit: int = 50) -> List[Dict]:
        """
        获取用户的所有讨论记录（按时间倒序）
        
        Args:
            user_id: 用户ID
            limit: 返回的最大记录数
            
        Returns:
            讨论记录列表
        """
        try:
            discussions = []
            pattern = f"{user_id}_*.json"
            json_dir = self.storage_base / "json"
            
            for filepath in json_dir.glob(pattern):
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        discussion_data = json.load(f)
                        discussions.append({
                            "filepath": str(filepath),
                            "metadata": discussion_data["metadata"],
                            "preview": self._generate_preview(discussion_data)
                        })
                except Exception as e:
                    logger.warning(f"加载文件失败 {filepath}: {e}")
                    continue
            
            # 按时间倒序排序
            discussions.sort(key=lambda x: x["metadata"]["timestamp"], reverse=True)
            return discussions[:limit]
            
        except Exception as e:
            logger.error(f"获取用户讨论记录失败: {e}")
            return []
    
    def export_to_docx(self, discussion_data: Dict, export_path: Optional[str] = None) -> str:
        """
        导出讨论记录为Word文档
        
        Args:
            discussion_data: 讨论数据
            export_path: 导出路径（可选）
            
        Returns:
            导出的文件路径
        """
        try:
            if export_path is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                export_path = self.storage_base / "exports" / f"discussion_{timestamp}.docx"
            
            doc = Document()
            
            # 添加标题
            title = doc.add_heading('临床多智能体讨论汇总报告', 0)
            title.alignment = 1  # 居中
            
            # 添加元数据表格
            self._add_metadata_table(doc, discussion_data)
            
            # 添加病历信息部分
            self._add_medical_context_section(doc, discussion_data)
            
            # 添加诊断评估部分
            self._add_diagnosis_section(doc, discussion_data)
            
            # 添加治疗方案部分
            self._add_treatment_section(doc, discussion_data)
            
            # 添加各科室意见整合
            self._add_specialty_integration_section(doc, discussion_data)
            
            # 添加随访计划
            self._add_followup_section(doc, discussion_data)
            
            # 添加讨论质量评估
            self._add_quality_assessment_section(doc, discussion_data)
            
            # 添加详细讨论过程（可选）
            self._add_detailed_discussion_section(doc, discussion_data)
            
            doc.save(export_path)
            logger.info(f"讨论记录已导出为Word文档: {export_path}")
            return str(export_path)
            
        except Exception as e:
            logger.error(f"导出Word文档失败: {e}")
            raise
    
    def export_to_html(self, discussion_data: Dict, export_path: Optional[str] = None) -> str:
        """
        导出讨论记录为HTML文件
        
        Args:
            discussion_data: 讨论数据
            export_path: 导出路径（可选）
            
        Returns:
            导出的文件路径
        """
        try:
            if export_path is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                export_path = self.storage_base / "exports" / f"discussion_{timestamp}.html"
            
            html_content = self._generate_html_content(discussion_data)
            
            with open(export_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"讨论记录已导出为HTML: {export_path}")
            return str(export_path)
            
        except Exception as e:
            logger.error(f"导出HTML失败: {e}")
            raise
    
    def delete_discussion(self, discussion_id: str, user_id: str) -> bool:
        """
        删除指定的讨论记录
        
        Args:
            discussion_id: 讨论ID
            user_id: 用户ID
            
        Returns:
            删除是否成功
        """
        try:
            pattern = f"{user_id}_{discussion_id}_*.json"
            json_dir = self.storage_base / "json"
            
            for filepath in json_dir.glob(pattern):
                # 创建备份
                backup_dir = self.storage_base / "backups" / datetime.now().strftime("%Y%m%d")
                backup_dir.mkdir(exist_ok=True)
                backup_path = backup_dir / filepath.name
                
                import shutil
                shutil.copy2(filepath, backup_path)
                
                # 删除原文件
                filepath.unlink()
                logger.info(f"讨论记录已删除并备份: {discussion_id}")
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"删除讨论记录失败: {e}")
            return False
    
    def _generate_preview(self, discussion_data: Dict) -> Dict:
        """生成讨论记录的预览信息"""
        summary = discussion_data.get("clinical_summary", {})
        return {
            "primary_diagnosis": summary.get("primary_diagnosis", "未知"),
            "diagnosis_confidence": summary.get("diagnosis_confidence", "未知"),
            "key_treatments": summary.get("key_treatments", [])[:3],
            "quality_score": summary.get("quality_score", 0)
        }
    
    def _add_metadata_table(self, doc: Document, data: Dict):
        """添加元数据表格"""
        doc.add_heading('报告信息', level=1)
        
        metadata = data["metadata"]
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        cells = table.rows[0].cells
        cells[0].text = "项目"
        cells[1].text = "内容"
        
        # 数据行
        rows_data = [
            ("讨论ID", metadata.get("discussion_id", "未知")),
            ("创建时间", metadata.get("created_at", "未知")),
            ("参与智能体", ", ".join(metadata.get("agents_used", []))),
            ("讨论轮数", str(metadata.get("rounds", 0)))
        ]
        
        for i, (label, value) in enumerate(rows_data, 1):
            cells = table.rows[i].cells
            cells[0].text = label
            cells[1].text = value
    
    def _add_medical_context_section(self, doc: Document, data: Dict):
        """添加病历信息部分"""
        doc.add_heading('病历信息', level=1)
        
        context = data["medical_context"]
        doc.add_heading('主诉与现病史', level=2)
        doc.add_paragraph(context.get("medical_record", "无"))
        
        doc.add_heading('讨论问题', level=2)
        doc.add_paragraph(context.get("question", "无"))
        
        additional_info = context.get("user_additional_info")
        if additional_info:
            doc.add_heading('用户补充信息', level=2)
            doc.add_paragraph(additional_info)
    
    def _add_diagnosis_section(self, doc: Document, data: Dict):
        """添加诊断评估部分"""
        doc.add_heading('诊断评估', level=1)
        
        summary = data["clinical_summary"]
        
        doc.add_heading('主要诊断', level=2)
        doc.add_paragraph(summary.get("primary_diagnosis", "未明确"))
        
        doc.add_heading('鉴别诊断', level=2)
        differential = summary.get("differential_diagnosis", [])
        if differential:
            for dd in differential:
                p = doc.add_paragraph(dd, style='List Bullet')
        else:
            doc.add_paragraph("无明确的鉴别诊断")
        
        doc.add_heading('诊断置信度', level=2)
        doc.add_paragraph(str(summary.get("diagnosis_confidence", "未知")))
    
    def _add_treatment_section(self, doc: Document, data: Dict):
        """添加治疗方案部分"""
        doc.add_heading('治疗方案', level=1)
        
        summary = data["clinical_summary"]
        treatments = summary.get("treatment_plan", {})
        
        for category, plan in treatments.items():
            doc.add_heading(category, level=2)
            if isinstance(plan, list):
                for item in plan:
                    doc.add_paragraph(item, style='List Bullet')
            else:
                doc.add_paragraph(str(plan))
    
    def _add_specialty_integration_section(self, doc: Document, data: Dict):
        """添加各科室意见整合"""
        doc.add_heading('各科室意见整合', level=1)
        
        summary = data["clinical_summary"]
        integration = summary.get("specialty_integration", {})
        
        for specialty, opinion in integration.items():
            doc.add_heading(specialty, level=2)
            doc.add_paragraph(opinion)
    
    def _add_followup_section(self, doc: Document, data: Dict):
        """添加随访计划"""
        doc.add_heading('随访计划', level=1)
        
        summary = data["clinical_summary"]
        followup = summary.get("follow_up_plan", [])
        
        if followup:
            for item in followup:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph("无具体的随访计划")
    
    def _add_quality_assessment_section(self, doc: Document, data: Dict):
        """添加讨论质量评估"""
        doc.add_heading('讨论质量评估', level=1)
        
        metrics = data.get("evaluation_metrics", {})
        
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        cells = table.rows[0].cells
        cells[0].text = "评估维度"
        cells[1].text = "评分/评价"
        
        assessment_items = [
            ("诊断全面性", metrics.get("diagnosis_completeness", "未评估")),
            ("治疗方案合理性", metrics.get("treatment_rationality", "未评估")),
            ("意见整合度", metrics.get("integration_quality", "未评估")),
            ("逻辑严谨性", metrics.get("logical_rigor", "未评估")),
            ("证据支持度", metrics.get("evidence_support", "未评估")),
            ("总体质量得分", f"{metrics.get('overall_score', 0)}/100")
        ]
        
        for i, (dimension, evaluation) in enumerate(assessment_items, 1):
            cells = table.rows[i].cells
            cells[0].text = dimension
            cells[1].text = str(evaluation)
    
    def _add_detailed_discussion_section(self, doc: Document, data: Dict):
        """添加详细讨论过程（可选）"""
        if len(data["discussion_process"]["discussion_log"]) > 10:  # 如果讨论记录太多，不详细显示
            doc.add_heading('详细讨论过程', level=1)
            doc.add_paragraph("讨论过程较为详细，如需查看完整记录请导出JSON格式文件。")
            return
        
        doc.add_heading('讨论过程摘要', level=1)
        discussion_log = data["discussion_process"]["discussion_log"]
        
        for i, round_data in enumerate(discussion_log):
            doc.add_heading(f'第{i+1}轮讨论', level=2)
            
            for contribution in round_data.get("contributions", []):
                agent = contribution.get("agent", "未知智能体")
                reasoning = contribution.get("contribution", {}).get("reasoning", "")
                
                doc.add_heading(agent, level=3)
                doc.add_paragraph(reasoning[:500] + "..." if len(reasoning) > 500 else reasoning)
    
    def _generate_html_content(self, data: Dict) -> str:
        """生成HTML内容"""
        summary = data["clinical_summary"]
        metrics = data.get("evaluation_metrics", {})
        
        html_template = """
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <title>临床多智能体讨论报告</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; }
                .header { text-align: center; border-bottom: 2px solid #333; padding-bottom: 20px; }
                .section { margin: 30px 0; }
                .subsection { margin: 20px 0; }
                table { border-collapse: collapse; width: 100%; margin: 10px 0; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; }
                .metric-score { font-weight: bold; color: #2c3e50; }
            </style>
        </head>
        <body>
            <div class="header">
                <h1>临床多智能体讨论汇总报告</h1>
                <p>生成时间: {timestamp}</p>
            </div>
            
            <div class="section">
                <h2>诊断评估</h2>
                <div class="subsection">
                    <h3>主要诊断</h3>
                    <p>{primary_diagnosis}</p>
                </div>
                <div class="subsection">
                    <h3>鉴别诊断</h3>
                    <ul>{differential_diagnosis}</ul>
                </div>
            </div>
            
            <div class="section">
                <h2>质量评估</h2>
                <table>
                    <tr><th>评估维度</th><th>评分</th></tr>
                    {metrics_table}
                </table>
            </div>
        </body>
        </html>
        """
        
        # 生成鉴别诊断列表
        dd_html = ""
        for dd in summary.get("differential_diagnosis", []):
            dd_html += f"<li>{html.escape(dd)}</li>"
        
        # 生成评估指标表格
        metrics_html = ""
        metrics_items = [
            ("诊断全面性", metrics.get("diagnosis_completeness", "未评估")),
            ("治疗方案合理性", metrics.get("treatment_rationality", "未评估")),
            ("意见整合度", metrics.get("integration_quality", "未评估")),
            ("总体质量得分", f"{metrics.get('overall_score', 0)}/100")
        ]
        
        for dimension, score in metrics_items:
            metrics_html += f"<tr><td>{dimension}</td><td class='metric-score'>{score}</td></tr>"
        
        return html_template.format(
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            primary_diagnosis=html.escape(summary.get("primary_diagnosis", "未明确")),
            differential_diagnosis=dd_html,
            metrics_table=metrics_html
        )

    def get_storage_stats(self, user_id: str) -> Dict:
        """获取存储统计信息"""
        user_discussions = self.get_user_discussions(user_id)
        total_size = 0
        
        for discussion in user_discussions:
            filepath = Path(discussion["filepath"])
            if filepath.exists():
                total_size += filepath.stat().st_size
        
        return {
            "total_discussions": len(user_discussions),
            "total_size_mb": round(total_size / (1024 * 1024), 2),
            "oldest_discussion": user_discussions[-1]["metadata"]["timestamp"] if user_discussions else None,
            "newest_discussion": user_discussions[0]["metadata"]["timestamp"] if user_discussions else None
        }
    def export_discussion(self, discussion_data: Dict, format: str, export_path: str = None) -> str:
        """
        导出讨论记录为指定格式 - 增强错误处理和重试逻辑
        """
        # === 修改：先验证导出路径和格式 ===
        if format not in ["docx", "html", "simple_html", "json", "txt"]:
            raise ValueError(f"不支持的导出格式: {format}")
        
        try:
            if format == "docx":
                return self.export_to_docx(discussion_data, export_path)
            elif format == "html":
                return self.export_to_html(discussion_data, export_path)
            elif format == "simple_html":
                return self.export_to_simple_html(discussion_data, export_path)
            elif format == "json":
                return self.export_to_json(discussion_data, export_path)
            elif format == "txt":
                return self.export_to_txt(discussion_data, export_path)
                
        except Exception as e:
            logger.error(f"导出失败: {e}")
            # 重新抛出异常，让调用者处理
            raise

    def export_to_json(self, discussion_data: Dict, export_path: str = None) -> str:
        """导出为JSON格式"""
        try:
            if export_path is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                export_path = self.storage_base / "exports" / f"discussion_{timestamp}.json"
            
            with open(export_path, 'w', encoding='utf-8') as f:
                json.dump(discussion_data, f, ensure_ascii=False, indent=2)
            
            logger.info(f"JSON导出完成: {export_path}")
            return str(export_path)
            
        except Exception as e:
            logger.error(f"JSON导出失败: {e}")
            raise    

    def export_to_simple_html(self, discussion_data: Dict, export_path: str = None) -> str:
        """
        导出为简化HTML格式 - 增加格式验证
        """
        # === 新增：验证数据格式 ===
        if not isinstance(discussion_data, dict):
            raise ValueError("讨论数据必须是字典格式")
            
        required_fields = ['metadata', 'medical_context', 'discussion_process']
        for field in required_fields:
            if field not in discussion_data:
                raise ValueError(f"缺少必要字段: {field}")
        
        try:
            if export_path is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                export_path = self.storage_base / "exports" / f"discussion_simple_{timestamp}.html"
            
            html_content = self._generate_simple_html_content(discussion_data)
            
            # === 新增：验证生成的HTML内容 ===
            if not html_content or len(html_content.strip()) < 100:
                raise ValueError("生成的HTML内容过短或为空")
                
            # 检查基本的HTML标签
            if "<html" not in html_content or "<body" not in html_content:
                raise ValueError("生成的HTML格式不正确")
            
            with open(export_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"简化HTML导出完成: {export_path}")
            return str(export_path)
            
        except Exception as e:
            logger.error(f"简化HTML导出失败: {e}")
            raise

    def _generate_simple_html_content(self, data: Dict) -> str:
        """生成简化的HTML内容 - 修复未定义字段问题"""
        
        # === 修复：安全地获取数据字段，提供默认值 ===
        # 安全获取metadata，提供默认值
        metadata = data.get('metadata', {})
        discussion_id = metadata.get('discussion_id', '未知ID')
        agents_used = metadata.get('agents_used', [])
        total_rounds = metadata.get('rounds', 0)
        
        # 安全获取medical_context，提供默认值
        medical_context = data.get('medical_context', {})
        question = medical_context.get('question', '无问题描述')
        
        # 安全获取discussion_log，提供默认值
        discussion_log = data.get('discussion_process', {}).get('discussion_log', [])
        
        # 安全获取clinical_summary，提供默认值
        clinical_summary = data.get('clinical_summary', {})
        if isinstance(clinical_summary, dict):
            final_decision = clinical_summary.get('final_decision', clinical_summary.get('summary', '暂无汇总意见'))
        else:
            final_decision = str(clinical_summary) if clinical_summary else '暂无汇总意见'

        # === 修复后的HTML模板 ===
        html_template = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>临床讨论简版报告</title>
    <style>
        body {{ 
            font-family: "Microsoft YaHei", Arial, sans-serif; 
            margin: 40px; 
            line-height: 1.6;
            color: #333;
        }}
        .header {{ 
            text-align: center; 
            border-bottom: 2px solid #2c3e50; 
            padding-bottom: 20px; 
            margin-bottom: 30px;
        }}
        .section {{ 
            margin: 30px 0; 
            padding: 20px;
            border-left: 4px solid #3498db;
            background-color: #f8f9fa;
        }}
        .round {{ 
            margin: 20px 0; 
            padding: 15px;
            border: 1px solid #ddd;
            border-radius: 5px;
            background-color: #fff;
        }}
        .agent-contribution {{ 
            margin: 15px 0; 
            padding: 10px;
            border-left: 3px solid #27ae60;
            background-color: #f1f8e9;
        }}
        .summary {{ 
            margin: 20px 0; 
            padding: 20px;
            border: 2px solid #e74c3c;
            border-radius: 5px;
            background-color: #fff5f5;
        }}
        .agent-name {{ 
            font-weight: bold; 
            color: #2c3e50;
            margin-bottom: 5px;
        }}
        .timestamp {{ 
            color: #7f8c8d; 
            font-size: 0.9em;
        }}
        h1 {{ color: #2c3e50; }}
        h2 {{ color: #34495e; border-bottom: 1px solid #ecf0f1; padding-bottom: 10px; }}
        h3 {{ color: #16a085; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>🏥🏥 临床多智能体讨论简版报告</h1>
        <p class="timestamp">生成时间: {timestamp} | 讨论ID: {discussion_id}</p>
    </div>
    
    <div class="section">
        <h2>📋📋 基本信息</h2>
        <p><strong>参与科室:</strong> {agents}</p>
        <p><strong>讨论轮数:</strong> {total_rounds}</p>
        <p><strong>讨论问题:</strong> {question}</p>
    </div>
    
    <div class="section">
        <h2>💬💬 讨论过程</h2>
        {rounds_content}
    </div>
    
    <div class="section summary">
        <h2>📊📊 最终汇总意见</h2>
        {summary_content}
    </div>
</body>
</html>
"""
        
        # === 修复：构建讨论轮次内容，处理空数据 ===
        rounds_content = ""
        if discussion_log:
            for round_data in discussion_log:
                round_num = round_data.get("round", 0)
                rounds_content += f'<div class="round">\n'
                rounds_content += f'<h3>第{round_num + 1}轮讨论</h3>\n'
                
                contributions = round_data.get("contributions", [])
                for contribution in contributions:
                    agent_name = contribution.get("agent", "未知科室")
                    
                    # 安全获取分析内容
                    contribution_data = contribution.get("contribution", {})
                    if isinstance(contribution_data, dict):
                        analysis = contribution_data.get("concise_analysis", 
                                                        contribution_data.get("analysis", 
                                                        contribution_data.get("response", "无分析内容")))
                    else:
                        analysis = str(contribution_data)
                    
                    timestamp = contribution.get("timestamp", "")
                    
                    rounds_content += f'''
                    <div class="agent-contribution">
                        <div class="agent-name">🏥🏥 {agent_name}</div>
                        <div class="timestamp">🕒🕒🕒 {timestamp}</div>
                        <div class="analysis">{analysis}</div>
                    </div>
                    '''
                
                rounds_content += '</div>\n'
        else:
            rounds_content = '<p>暂无讨论记录</p>'
        
        # === 修复：构建汇总内容，处理各种可能的数据格式 ===
        summary_content = ""
        if isinstance(clinical_summary, dict):
            if clinical_summary.get("final_decision"):
                summary_content = f'<div class="final-decision">{clinical_summary["final_decision"]}</div>'
            elif clinical_summary.get("summary"):
                summary_content = f'<div class="final-decision">{clinical_summary["summary"]}</div>'
            else:
                # 尝试从其他字段获取汇总信息
                for key in ["diagnosis", "treatment_plan", "conclusion"]:
                    if key in clinical_summary:
                        summary_content = f'<div class="final-decision">{clinical_summary[key]}</div>'
                        break
                else:
                    summary_content = "<p>暂无汇总意见</p>"
        else:
            summary_content = f'<div class="final-decision">{str(clinical_summary)}</div>'
        
        return html_template.format(
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            discussion_id=discussion_id,
            agents=", ".join(agents_used) if agents_used else "无参与科室",
            total_rounds=total_rounds,
            question=question,
            rounds_content=rounds_content,
            summary_content=summary_content
        )

    def save_discussion(self, user_id: str, discussion_data: Dict) -> str:
        """
        保存讨论记录到JSON文件 - 确保数据结构完整
        """
        try:
            # 生成唯一讨论ID和时间戳
            discussion_id = str(uuid.uuid4())[:8]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{user_id}_{discussion_id}_{timestamp}.json"
            filepath = self.storage_base / "json" / filename
            
            # === 修复：确保数据结构包含所有必要字段 ===
            # 构建完整的metadata
            metadata = discussion_data.get("metadata", {})
            if not metadata:
                metadata = {
                    "discussion_id": discussion_id,
                    "user_id": user_id,
                    "timestamp": timestamp,
                    "created_at": datetime.now().isoformat(),
                    "agents_used": discussion_data.get("agents", []),
                    "rounds": discussion_data.get("rounds", 0),
                    "medical_record_length": len(discussion_data.get("medical_record", "")),
                    "question_length": len(discussion_data.get("question", ""))
                }
            
            # 构建完整的medical_context
            medical_context = discussion_data.get("medical_context", {})
            if not medical_context:
                medical_context = {
                    "medical_record": discussion_data.get("medical_record", ""),
                    "question": discussion_data.get("question", ""),
                    "user_additional_info": discussion_data.get("user_additional_info", "")
                }
            
            # 构建完整的discussion_process
            discussion_process = discussion_data.get("discussion_process", {})
            if not discussion_process:
                discussion_process = {
                    "discussion_log": discussion_data.get("log", []),
                    "user_interventions": discussion_data.get("interventions", []),
                    "logic_reports": discussion_data.get("logic_reports", [])
                }
            
            # 构建完整的讨论记录
            discussion_record = {
                "metadata": metadata,
                "medical_context": medical_context,
                "discussion_process": discussion_process,
                "clinical_summary": discussion_data.get("summary", {}),
                "evaluation_metrics": discussion_data.get("metrics", {})
            }
            
            # 保存到JSON文件
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(discussion_record, f, ensure_ascii=False, indent=2)
            
            logger.info(f"讨论记录已保存: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"保存讨论记录失败: {e}")
            raise

    def load_discussion(self, discussion_id: str, user_id: str) -> Optional[Dict]:
        """
        加载特定的讨论记录 - 提供默认值
        """
        try:
            # 查找匹配的文件
            pattern = f"{user_id}_{discussion_id}_*.json"
            json_dir = self.storage_base / "json"
            
            for filepath in json_dir.glob(pattern):
                with open(filepath, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # === 修复：确保返回的数据结构完整 ===
                return {
                    "metadata": data.get("metadata", {}),
                    "medical_context": data.get("medical_context", {}),
                    "discussion_process": data.get("discussion_process", {}),
                    "clinical_summary": data.get("clinical_summary", {}),
                    "evaluation_metrics": data.get("evaluation_metrics", {})
                }
            
            logger.warning(f"未找到讨论记录: {discussion_id}")
            return None
            
        except Exception as e:
            logger.error(f"加载讨论记录失败: {e}")
            return None


